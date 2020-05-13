'''
create *.docx file based on PDF layout data with python package python-docx.
@created: 2019-06-28
@author: train8808@gmail.com
'''


from io import BytesIO

from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from . import utils



def make_page(doc, layout):
    ''' create page based on layout data. 

        To avoid incorrect page break from original document, a new page section
        is created for each page.

        Support general document style only:
          - writing mode: from left to right, top to bottom
          - text direction: horizontal

        The vertical postion of paragraph/table is defined by space_before or 
        space_after property of a paragraph.
    '''
    # new page section
    # a default section is created when initialize the document,
    # so we do not have to add section for the first time.
    if not doc.paragraphs:
        section = doc.sections[0]
    else:
        section = doc.add_section(WD_SECTION.NEW_PAGE)

    width, height = layout['width'], layout['height']
    section.page_width  = Pt(width)
    section.page_height = Pt(height)

    # set page margin
    left,right,top,bottom = layout['margin']    
    section.left_margin = Pt(left)
    section.right_margin = Pt(right)
    section.top_margin = Pt(top)
    section.bottom_margin = Pt(bottom)

    # add paragraph or table according to parsed block
    for block in layout['blocks']:
        # make paragraphs
        if block['type'] in (0, 1):
            # horizontal paragraph
            if block['type']==1 or block['lines'][0]['wmode'] == 0:
                make_paragraph(doc, block, width, layout['margin'])
            
            # vertical paragraph
            else:
                make_vertical_paragraph(doc, block)
        
        # make table
        elif block['type']==3:
            make_table(doc, block, width, layout['margin'])            


def make_paragraph(doc, block, width, page_margin):
    '''create paragraph for a text block.
       join line sets with TAB and set position according to bbox.

       Generally, a pdf block is a docx paragraph, with block|line as line in paragraph.
       But without the context, it's not able to recognize a block line as word wrap, or a 
       separate line instead. A rough rule used here:
        - block line will be treated as separate line (append `\n`) by default, except
        - (1) this line and next line are actually in the same line (y-position)
        - (2) if the rest space of this line can't accommodate even one span of next line, 
              it's supposed to be normal word wrap.
    '''
    # new paragraph    
    p = doc.add_paragraph()

    # indent and space setting
    before_spacing = max(round(block.get('before_space', 0.0), 1), 0.0)
    after_spacing = max(round(block.get('after_space', 0.0), 1), 0.0)
    pf = reset_paragraph_format(p)
    pf.space_before = Pt(before_spacing)
    pf.space_after = Pt(after_spacing)    

    # add image
    if block['type']==1:
        # left indent implemented with tab
        pos = block['bbox'][0]-page_margin[0]
        if abs(pos) > utils.DM:
            pf.tab_stops.add_tab_stop(Pt(pos))
            p.add_run().add_tab()
        # create image with bytes data stored in block.
        span = p.add_run()
        span.add_picture(BytesIO(block['image']), width=Pt(block['bbox'][2]-block['bbox'][0]))

    # add text (inline image may exist)
    else:
        # set line spacing for text paragraph
        pf.line_spacing = Pt(round(block['line_space'],1))

        for i, line in enumerate(block['lines']):

            # left indent implemented with tab
            pos = line['bbox'][0]-page_margin[0]
            if abs(pos) > utils.DM:
                pf.tab_stops.add_tab_stop(Pt(pos))
                p.add_run().add_tab()

            # add line
            for span in line['spans']:
                # add content
                add_span(span, p)

                # exactly line spacing will destroy image display, so set single line spacing instead
                if 'image' in span:
                    pf.line_spacing = 1.05

            # break line or word wrap?
            # new line by default
            line_break = True

            # no more lines after last line
            if line==block['lines'][-1]: 
                line_break = False
            
            # different lines in space, i.e. break line if they are not horizontally aligned
            # Line i+1 y0 > Line i y1 is a simple criterion, but not so general since overlap may exist
            # so a overlap with at least 0.5 times of line width is applied here
            elif utils.is_horizontal_aligned(block['lines'][i+1]['bbox'], line['bbox'], True, 0.5):
                line_break = False
            
            # now, we have two lines, check whether word wrap or line break
            else:
                # bbox of first span in next line
                x0, _, x1, _ = block['lines'][i+1]['spans'][0]['bbox']
                # word wrap if rest space of this line can't accommodate
                # even one span of next line
                free_space = width-page_margin[1]-line['bbox'][2]
                if x1-x0 >= free_space:
                    line_break = False
            
            if line_break:
                p.add_run('\n')

    return p
    

def make_vertical_paragraph(doc, block):
    pass


def make_table(doc, block, page_width, page_margin):
    '''create table for a text block
       count of columns are checked, combine rows if next block is also in table format
    '''
    # new table
    block_cells = block['cells']
    table = doc.add_table(rows=len(block_cells), cols=len(block_cells[0]))

    # set indent
    pos = block['bbox'][0]-page_margin[0]
    indent_table(table, pos)

    # cell format and contents
    for i, (row, block_row) in enumerate(zip(table.rows, block_cells)):
        for j, (cell, block_cell) in enumerate(zip(row.cells, block_row)):

            # ignore merged cells
            if not block_cell: continue

            # merge cells
            n_row, n_col = block_cell['merged-cells']
            if n_row*n_col!=1:
                _cell = table.cell(i+n_row-1, j+n_col-1)
                cell.merge(_cell)

            # set borders
            keys = ('top', 'end', 'bottom', 'left')
            kwargs = {}
            for k, w, c in zip(keys, block_cell['border-width'], block_cell['border-color']):
                hex_c = f'#{hex(c)[2:].zfill(6)}'
                kwargs[k] = {
                    'sz': w, 'val': 'single', 'color': hex_c.upper()
                }
            # merged cells should also be considered
            for m in range(i, i+n_row):
                for n in range(j, j+n_col):
                    set_cell_border(table.cell(m, n), **kwargs)

            # set width/height

            # set bg-color
            if block_cell['bg-color']!=None:
                set_cell_shading(cell, block_cell['bg-color'])

            # insert text
            p = cell.paragraphs[0]
            for line in block_cell['lines']:
                # add line
                for span in line['spans']:
                    # add content
                    add_span(span, p)


def reset_paragraph_format(p):
    '''paragraph format'''
    pf = p.paragraph_format
    pf.line_spacing = 1.05 # single
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.widow_control = True
    return pf


def add_span(span, paragraph):
    '''add text span to a paragraph.       
    '''
    # inline image span
    if 'image' in span:
        image_span = paragraph.add_run()
        image_span.add_picture(BytesIO(span['image']), width=Pt(span['bbox'][2]-span['bbox'][0]))

    # text span
    else:
        text_span = paragraph.add_run(span['text'])

        # style setting
        # https://python-docx.readthedocs.io/en/latest/api/text.html#docx.text.run.Font

        # basic font style
        # line['flags'] is an integer, encoding bool of font properties:
        # bit 0: superscripted (2^0)
        # bit 1: italic (2^1)
        # bit 2: serifed (2^2)
        # bit 3: monospaced (2^3)
        # bit 4: bold (2^4)            
        text_span.italic = bool(span['flags'] & 2**1)
        text_span.bold = bool(span['flags'] & 2**4)
        text_span.font.name = utils.parse_font_name(span['font'])
        text_span.font.size = Pt(round(span['size']*2)/2.0) # only x.0 and x.5 is accepted in docx
        text_span.font.color.rgb = RGBColor(*utils.RGB_component(span['color']))

        # font style parsed from PDF rectangles: 
        # e.g. highlight, underline, strike-through-line
        for style in span.get('style', []):
            t = style['type']
            if t==0:
                text_span.font.highlight_color = utils.to_Highlight_color(style['color'])
            elif t==1:
                text_span.font.underline = True
            elif t==2:
                text_span.font.strike = True



def indent_table(table, indent):
    '''indent table

       args:
         - indent: indent value, the basic unit is 1/20 pt
    '''
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(20*indent)) # basic unit 1/20 pt for openxml 
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)


def set_cell_shading(cell, RGB_value):
    '''set cell background-color'''
    c = hex(RGB_value)[2:].zfill(6)
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), c)))


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border, refer to:
    https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx

    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))